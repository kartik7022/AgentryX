from __future__ import annotations

from pathlib import Path


def create_pdf(path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError as exc:  # pragma: no cover
        raise SystemExit(f"reportlab is required to generate PDF fixtures: {exc}")

    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawString(72, 720, "DocAI sample PDF")
    c.drawString(72, 700, "Invoice INV-2026-001 for ₹45,000 from ABC Tech")
    c.save()


def create_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("DocAI Sample DOCX", level=1)
    doc.add_paragraph("This is a sample document for parser tests.")
    doc.add_paragraph("Purchase order PO #4521 pending approval from finance team.")
    doc.save(str(path))


def create_png(path: Path) -> None:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (1200, 800), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 50), "DocAI sample PNG", fill="black")
    draw.text((50, 100), "Passport scan for Vinod Kumar, expiry 2027-05-01", fill="black")
    image.save(str(path))


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    fixtures = root / "tests" / "fixtures"
    fixtures.mkdir(parents=True, exist_ok=True)

    create_pdf(fixtures / "sample.pdf")
    create_docx(fixtures / "sample.docx")
    create_png(fixtures / "sample.png")


if __name__ == "__main__":
    main()
